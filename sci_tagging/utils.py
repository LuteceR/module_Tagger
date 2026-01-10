import re
import zipfile
import xml.etree.ElementTree as ET
import hashlib
import logging
from typing import Tuple, List, Dict, Optional

LOG = logging.getLogger(__name__)

def clean_text(t: str) -> str:
    t = t.replace("\r", " ")
    t = re.sub(r"\s+", " ", t)
    return t.strip()

def extract_tags1(predicted_text: str) -> List[str]:
    pairs = re.findall(r"(?s)\[([A-ZА-ЯЁ0-9_]+)\*\](.*?)\[\*\1\]", predicted_text)
    tags = []
    for _, inner in pairs:
        t = " ".join(inner.split())
        if t:
            tags.append(t)
    tags = [t for t in tags if len(t) >= 2]
    return sorted(set(tags))

def read_docx_paragraphs(path: str) -> List[str]:
    try:
        with zipfile.ZipFile(path) as zf:
            with zf.open("word/document.xml") as f:
                xml_bytes = f.read()
        root = ET.fromstring(xml_bytes)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        for parent in root.iter():
            for child in list(parent):
                if child.tag == f"{{{ns['w']}}}tbl":
                    parent.remove(child)

        paragraphs = []
        for p in root.findall('.//w:p', ns):
            runs_text = []
            for t in p.findall('.//w:t', ns):
                runs_text.append(t.text or "")
            if runs_text:
                paragraphs.append("".join(runs_text))
        return paragraphs
    except Exception:
        LOG.info("XML-разбор не удался, пробую python-docx")

    try:
        from docx import Document  # type: ignore
        doc = Document(path)
        return [p.text for p in doc.paragraphs]
    except Exception:
        return []

def read_docx_first_page_text(path: str) -> str:
    try:
        with zipfile.ZipFile(path) as zf:
            with zf.open("word/document.xml") as f:
                xml_bytes = f.read()
        root = ET.fromstring(xml_bytes)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        body = root.find('w:body', ns)
        if body is None:
            return ""

        collected = []
        for elem in body.iter():
            if elem.tag == f"{{{ns['w']}}}br" and elem.attrib.get(f"{{{ns['w']}}}type") == "page":
                break
            if elem.tag == f"{{{ns['w']}}}lastRenderedPageBreak":
                break
            if elem.tag == f"{{{ns['w']}}}p":
                runs_text = []
                for t in elem.findall('.//w:t', ns):
                    runs_text.append(t.text or "")
                if runs_text:
                    collected.append("".join(runs_text))
        return "\n".join(collected)
    except Exception:
        paras = read_docx_paragraphs(path)
        return "\n".join(paras[:20])

def _normalize_heading_text(s: str) -> str:
    return (s or "").strip().lower().replace("ё", "е")

def _find_main_range_via_toc(path: str) -> Tuple[Optional[int], Optional[int]]:
    try:
        with zipfile.ZipFile(path) as zf:
            with zf.open("word/document.xml") as f:
                xml_bytes = f.read()
        root = ET.fromstring(xml_bytes)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        for parent in root.iter():
            for child in list(parent):
                if child.tag == f"{{{ns['w']}}}tbl":
                    parent.remove(child)

        paragraphs = []
        bookmark_to_pindex: Dict[str, int] = {}
        for p_idx, p in enumerate(root.findall('.//w:p', ns)):
            paragraphs.append(p)
            for bmk in p.findall('.//w:bookmarkStart', ns):
                name = bmk.attrib.get(f"{{{ns['w']}}}name", "")
                if name:
                    bookmark_to_pindex[name] = p_idx

        start_anchor = None
        end_anchor = None
        for hyperlink in root.findall('.//w:hyperlink', ns):
            anchor = hyperlink.attrib.get(f"{{{ns['w']}}}anchor")
            if not anchor:
                continue
            text_runs = [t.text or "" for t in hyperlink.findall('.//w:t', ns)]
            link_text = _normalize_heading_text("".join(text_runs))
            if not link_text:
                continue
            if link_text.startswith("введение"):
                start_anchor = anchor
            if link_text.startswith("заключение") or link_text.startswith("вывод"):
                end_anchor = anchor

        start_idx = bookmark_to_pindex.get(start_anchor) if start_anchor else None
        end_idx = bookmark_to_pindex.get(end_anchor) if end_anchor else None
        return (start_idx, end_idx)
    except Exception:
        return (None, None)

def _find_main_range_via_styles(path: str, paragraphs_text: List[str]) -> Tuple[Optional[int], Optional[int]]:
    try:
        from docx import Document  # type: ignore
        doc = Document(path)
        start_idx = None
        end_idx = None
        for idx, p in enumerate(doc.paragraphs):
            name = (getattr(p.style, 'name', '') or '').lower()
            if not name:
                continue
            if not ("heading" in name or "заголов" in name):
                continue
            txt = _normalize_heading_text(p.text)
            if not txt:
                continue
            if start_idx is None and txt.startswith("введение"):
                start_idx = idx
            if end_idx is None and (txt.startswith("заключение") or txt.startswith("вывод")):
                end_idx = idx
        return (start_idx, end_idx)
    except Exception:
        return (None, None)

def _find_main_range_via_regex(paragraphs_text: List[str]) -> Tuple[Optional[int], Optional[int]]:
    start_idx = None
    end_idx = None
    start_re = re.compile(r"^(\d+\.?\s*)?введение\b", re.IGNORECASE)
    end_re = re.compile(r"^(\d+\.?\s*)?(заключение|выводы?)\b", re.IGNORECASE)
    for i, t in enumerate(paragraphs_text):
        txt = _normalize_heading_text(t)
        if start_idx is None and start_re.search(txt):
            start_idx = i
        if end_idx is None and end_re.search(txt):
            end_idx = i
    return (start_idx, end_idx)

def select_main_paragraph_range(path: str, paragraphs_text: List[str]) -> Tuple[int, int]:
    s_idx, e_idx = _find_main_range_via_toc(path)
    if s_idx is None and e_idx is None:
        s_idx, e_idx = _find_main_range_via_styles(path, paragraphs_text)
    if s_idx is None and e_idx is None:
        s_idx, e_idx = _find_main_range_via_regex(paragraphs_text)
    n = len(paragraphs_text)
    if s_idx is None and e_idx is None:
        return (0, max(0, n - 1))
    if s_idx is None:
        return (0, min(e_idx, n - 1))
    if e_idx is None:
        return (max(0, s_idx), max(0, n - 1))
    if s_idx > e_idx:
        s_idx, e_idx = e_idx, s_idx
    return (max(0, s_idx), min(e_idx, n - 1))

def split_into_sentences(text: str) -> List[str]:
    sents = re.split(r"[.!?]+", text)
    return [s.strip() for s in sents if s and s.strip()]

def build_chunks_by_sentences(text: str, sentences_per_chunk: int) -> List[str]:
    sents = split_into_sentences(text)
    if sentences_per_chunk <= 0:
        sentences_per_chunk = 4
    chunks = [". ".join(sents[i:i + sentences_per_chunk]) for i in range(0, len(sents), sentences_per_chunk)]
    return [c for c in chunks if c]

def build_chunks_by_paragraphs(paragraphs: List[str], paragraphs_per_chunk: int) -> List[str]:
    paras = [p.strip() for p in paragraphs if p and p.strip()]
    if paragraphs_per_chunk <= 0:
        paragraphs_per_chunk = 1
    chunks = ["\n".join(paras[i:i + paragraphs_per_chunk]) for i in range(0, len(paras), paragraphs_per_chunk)]
    return [c for c in chunks if c]

def build_chunks_by_characters(text: str, chunk_chars: int, stride_chars: Optional[int] = None) -> List[str]:
    t = text.strip()
    if chunk_chars <= 0:
        chunk_chars = 2000
    if not stride_chars or stride_chars <= 0:
        stride_chars = chunk_chars
    chunks = []
    i = 0
    n = len(t)
    while i < n:
        chunks.append(t[i:i + chunk_chars])
        if i + chunk_chars >= n:
            break
        i += stride_chars
    return [c for c in chunks if c]

def sha256_file(path: str) -> str:
    try:
        h = hashlib.sha256()
        with open(path, "rb") as fbin:
            for chunk in iter(lambda: fbin.read(1 << 20), b""):
                h.update(chunk)
        return h.hexdigest()
    except Exception:
        return ""


